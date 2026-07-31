import os
import time
import threading
import schedule
import base64
import shutil
import uuid
import gradio as gr
from supabase import create_client, Client
from duckduckgo_search import DDGS

# 載入 Google GenAI 與 文書處理套件
try:
    import google.genai as genai
except ImportError:
    import genai

try:
    import docx
except ImportError:
    pass

try:
    from fpdf import FPDF
except ImportError:
    pass

# --- 1. 讀取雲端金鑰與多用戶設定 ---
SUPABASE_URL = os.environ.get("SUPABASE_URL")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY")
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")

users_env = os.environ.get("WEB_USERS", "admin:admin")
AUTH_LIST = []
if users_env:
    for pair in users_env.split(','):
        if ':' in pair:
            parts = pair.split(':', 1)
            if len(parts) == 2:
                username = parts[0].strip()
                password = parts[1].strip()
                if username and password:
                    AUTH_LIST.append((username, password))

supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY) if SUPABASE_URL and SUPABASE_KEY else None
client = genai.Client(api_key=GEMINI_API_KEY) if GEMINI_API_KEY else None

# --- 2. 雲端資料庫操作 ---
def get_user_profile(username):
    if not supabase: return "無"
    try:
        res = supabase.table('user_profile').select('fact').eq('username', username).execute()
        facts = [item['fact'] for item in res.data]
        return "\n".join([f"- {f}" for f in facts]) if facts else "無"
    except: return "無"

def update_user_profile(user_input, username):
    if not user_input or not supabase or not client: return
    try:
        draw_keywords = ["畫", "圖", "產出", "生成", "生一張", "存成", "輸出", "文件", "word", "doc", "pdf"]
        if any(k in user_input.lower() for k in draw_keywords):
            return

        prompt = f"分析這句話：「{user_input}」。是否透露了說話者的個人偏好、物品 or 習慣？有則總結事實，無則回覆『無』。"
        resp = client.models.generate_content(model='gemini-3.6-flash', contents=prompt)
        fact = resp.text.strip()
        if fact != "無" and "沒有" not in fact and len(fact) > 2:
            supabase.table('user_profile').insert({"fact": fact, "username": username}).execute()
    except: pass

def recall_long_term_memory(username):
    if not supabase: return []
    try:
        res = supabase.table('memory').select('*').eq('username', username).order('id', desc=True).limit(30).execute()
        return res.data
    except: return []

def search_the_web(query):
    draw_keywords = ["畫", "圖", "產出", "生成", "生一張", "存成", "文件", "word", "doc", "pdf"]
    if any(k in query.lower() for k in draw_keywords):
        return "特殊功能模式，不進行網路搜尋。"
    if not query: return "無文字查詢"
    try:
        results = DDGS().text(query, max_results=3)
        return "".join([f"【來源】{res['title']}\n摘要: {res['body']}\n" for res in results])
    except Exception:
        return "網路上暫時無法搜尋。"

# --- 3. 核心大腦邏輯 ---
def ask_smart_agent(user_text, uploaded_files, history, username):
    if not client or not supabase:
        return "⚠️ 系統尚未設定完整的 API 金鑰 (GEMINI 或 SUPABASE)，請至 Settings 中設定。"

    user_lower = user_text.lower()
    
    # --- V31 終極複合式訊息與 PDF 完美斷行修復 ---
    def create_file_response(content, file_type):
        doc_prompt = f"請根據以下指示，撰寫一份完整、專業的文章或報告內容（不需打招呼，直接給出純內文即可）：\n{content}"
        # ✅ 使用最新 3.6 模型
        resp = client.models.generate_content(model='gemini-3.6-flash', contents=doc_prompt)
        report_content = resp.text

        if file_type == 'pdf':
            font_path = "NotoSansTC-Regular.ttf"
            if not os.path.exists(font_path):
                return "❌ 找不到中文字型檔 (NotoSansTC-Regular.ttf)！請確認您已將字型上傳至 GitHub。"
            
            pdf = FPDF()
            pdf.add_page()
            pdf.add_font("NotoSansTC", "", font_path, uni=True)
            
            # 加入大標題
            pdf.set_font("NotoSansTC", size=18)
            pdf.cell(0, 15, txt="AI 智能分析報告", ln=1, align="C")
            
            # 寫入內文 (✅ 完美中文強制斷行演算法)
            pdf.set_font("NotoSansTC", size=12)
            max_w = pdf.w - 2 * pdf.l_margin - 5
            for line in report_content.split('\n'):
                if not line.strip():
                    pdf.cell(0, 10, txt="", ln=1)
                    continue
                current_line = ""
                for char in line:
                    if pdf.get_string_width(current_line + char) > max_w:
                        pdf.cell(0, 10, txt=current_line, ln=1)
                        current_line = char
                    else:
                        current_line += char
                if current_line:
                    pdf.cell(0, 10, txt=current_line, ln=1)
            
            filename = f"AI_Report_{uuid.uuid4().hex[:6]}.pdf"
            pdf.output(filename)
            
        elif file_type == 'docx':
            doc = docx.Document()
            doc.add_heading('AI 智能分析報告', 0)
            doc.add_paragraph(report_content)
            filename = f"AI_Report_{uuid.uuid4().hex[:6]}.docx"
            doc.save(filename)
        
        else:
            return "❌ 不支援的檔案類型。"
        
        return {
            "text": f"📝 您的 {file_type.upper()} 報告已生成！",
            "files": [filename]
        }

    # --- 文件產出攔截器 ---
    if "存成pdf" in user_lower or "產出pdf" in user_lower or "做成pdf" in user_lower:
        try:
            return create_file_response(user_text, 'pdf')
        except Exception as e:
            return f"❌ 產生 PDF 失敗：{str(e)}"
            
    doc_keywords = ["存成doc", "存成word", "產出文件", "輸出報告", "存成檔案", "做成word", "產出doc"]
    if any(k in user_lower for k in doc_keywords):
        try:
            return create_file_response(user_text, 'docx')
        except Exception as e:
            return f"❌ 產生 DOCX 失敗：{str(e)}"

    # ----- 一般文字與檔案邏輯 -----
    cleaned_history = []
    for item in history[-3:]:
        if isinstance(item, (list, tuple)) and len(item) >= 2:
            cleaned_history.append((item[0], item[1]))

    short_term = ""
    for user_msg, ai_msg in cleaned_history:
        if isinstance(user_msg, (list, tuple)): u_text = "[上傳了檔案]"
        elif isinstance(user_msg, dict): u_text = user_msg.get("text", "[上傳了檔案]")
        else: u_text = str(user_msg)
        safe_ai_msg = str(ai_msg)
        if "data:image" in safe_ai_msg: safe_ai_msg = "🎨 [AI 生成了一張圖片]"
        elif ".doc" in safe_ai_msg or ".pdf" in safe_ai_msg or "AI_Report" in safe_ai_msg: safe_ai_msg = "📝 [AI 提供了一份檔案]"
        short_term += f"我:{u_text}\n你:{safe_ai_msg}\n"
        
    profile = get_user_profile(username)
    long_term_mems = recall_long_term_memory(username)
    long_term_context = "".join([f"舊紀錄:{m['question']} -> {m['answer'][:100]}...\n" for m in long_term_mems])
    web_context = search_the_web(user_text)
    
    prompt = f"""你是一個多模態專屬 AI 助理。當前服務的主人是：{username}。
【主人長期特徵】\n{profile}
【本次對話上下文】\n{short_term if short_term else "新對話。"}
【歷史記憶(請自動判斷關聯性)】\n{long_term_context if long_term_context else "無。"}
【網路最新資訊】\n{web_context}
【目前提問/指示】\n{user_text}
請給出精準回答："""
    
    contents_to_send = [prompt]
    uploaded_g_files = [] 
    
    if uploaded_files:
        for file_info in uploaded_files:
            actual_path = file_info["path"] if isinstance(file_info, dict) else file_info
            file_size_mb = os.path.getsize(actual_path) / (1024 * 1024)
            if file_size_mb > 10: return f"⚠️ 警告：檔案 ({file_size_mb:.1f}MB) 超過 10MB 限制！"
            
            _, safe_ext = os.path.splitext(actual_path)
            safe_name = f"temp_upload_{uuid.uuid4().hex}{safe_ext}"
            try:
                shutil.copy(actual_path, safe_name)
                g_file = client.files.upload(file=safe_name)
                contents_to_send.append(g_file)
                uploaded_g_files.append(g_file) 
            except Exception as e:
                return f"❌ 檔案上傳失敗: {e}"
            finally:
                if os.path.exists(safe_name): os.remove(safe_name)

    try:
        response = client.models.generate_content(model='gemini-3.6-flash', contents=contents_to_send)
        final_answer = response.text
        db_question = user_text if user_text else "[分析了上傳的檔案]"
        supabase.table('memory').insert({"question": db_question, "answer": final_answer, "username": username}).execute()
        update_user_profile(user_text, username)
        return final_answer
    except Exception as e:
        return f"⚠️ AI 大腦暫時無法思考 (可能為安全攔截、金鑰無效或 API 限制)：{str(e)}"
    finally:
        for gf in uploaded_g_files:
            try: client.files.delete(name=gf.name)
            except: pass

# --- 4. 建立 Gradio 多模態網頁 ---
def chat_logic(message_dict, history, request: gr.Request):
    username = request.username if request and hasattr(request, "username") and request.username else "admin"
    raw_text = message_dict.get("text", "")
    text = raw_text.strip()
    files = message_dict.get("files", [])
    
    normalized_text = text.replace(" ", "").replace("　", "").replace("。", "").replace(".", "").replace("！", "").replace("!", "")
    if not text and files: text = "請幫我分析我上傳的檔案/錄音檔。"
    yield "⏳ 系統正在處理中，請稍候..."
    
    if not supabase:
        yield "⚠️ 錯誤：無法連線至 Supabase 資料庫，請檢查金鑰設定。"
        return
        
    if normalized_text == "清除所有對話紀錄":
        supabase.table('memory').delete().eq('username', username).execute()
        yield f"🧹 [{username}] 的所有歷史對話記憶已被永久刪除！"
    elif normalized_text == "清除我的個人畫像":
        supabase.table('user_profile').delete().eq('username', username).execute()
        yield f"🧹 [{username}] 的長期個人偏好筆記已被徹底清空！"
    elif normalized_text.startswith("設定目標：") or normalized_text.startswith("設定目標:"):
        target = raw_text.split("：")[-1].split(":")[-1].strip()
        supabase.table('research_targets').insert({"target": target, "username": username}).execute()
        yield f"🎯 [{username}] 的目標已設定：『{target}』"
    else:
        reply = ask_smart_agent(text, files, history, username)
        if isinstance(reply, dict) and "files" in reply:
            yield "📝 報告內容已撰寫完成，正在為您打包成檔案..."
            time.sleep(0.5)
            yield reply
        else:
            yield reply

demo = gr.ChatInterface(
    fn=chat_logic,
    multimodal=True,
    title="🚀 可進化 AI 助理 V31 (PDF 完美排版版)",
    description="具備多用戶隔離與 RAG 記憶的頂級架構。<br>👇 **請手動輸入，或點擊下方的【快捷指令按鈕】：**",
    examples=[
        [{"text": "幫我規劃三天兩夜的高雄旅遊行程，並存成PDF"}],
        [{"text": "幫我寫一份中秋節烤肉採買清單，並存成Word檔案"}],
        [{"text": "清除所有對話紀錄"}],
        [{"text": "設定目標：2026年人工智慧發展趨勢"}]
    ]
)

# ⚠️ 自動背景任務暫時關閉
# def daily_background_learning(): ...
# def run_scheduler(): ...
# threading.Thread(target=run_scheduler, daemon=True).start()

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 7860))
    if AUTH_LIST:
        demo.launch(server_name="0.0.0.0", server_port=port, auth=AUTH_LIST)
    else:
        demo.launch(server_name="0.0.0.0", server_port=port)
